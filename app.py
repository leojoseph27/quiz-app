import os
from flask import Flask, render_template, request, redirect, url_for, send_file, jsonify, session
import pdfplumber
import docx
from werkzeug.utils import secure_filename
import google.generativeai as genai
from fpdf import FPDF
from dotenv import load_dotenv
import time
import json  # ✅ ADDED
import uuid   # ✅ ADDED
from docx import Document as DocxDocument
from docx.shared import Pt

load_dotenv()

GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables")

genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel("gemini-2.5-flash")

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['RESULTS_FOLDER'] = 'results/'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'txt', 'doc', 'docx'}
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'your_secret_key')

# ✅ ADD QUIZ STORAGE DIRECTORY
QUIZ_STORAGE_DIR = 'quiz_storage/'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['RESULTS_FOLDER'], exist_ok=True)
os.makedirs(QUIZ_STORAGE_DIR, exist_ok=True)  # ✅ CREATE STORAGE FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extract_text_from_file(file_path):
    ext = file_path.rsplit('.', 1)[1].lower()
    try:
        print(f"Attempting to extract text from {file_path} with extension {ext}")
        if ext == 'pdf':
            text = ''
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                print(f"📄 Extracting {total_pages} pages...")
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
                    if (i + 1) % 10 == 0 or i == total_pages - 1:
                        print(f"   ✅ Processed page {i+1}/{total_pages}")
            print(f"✅ Successfully extracted {len(text)} characters from PDF")
            return text
        elif ext in ['doc', 'docx']:
            try:
                if not os.path.exists(file_path):
                    print(f"Word file not found at path: {file_path}")
                    return None

                doc = docx.Document(file_path)
                text = ''

                for para in doc.paragraphs:
                    if para.text.strip():
                        text += para.text + '\n'

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text += cell.text + '\n'

                if not text.strip():
                    print("No text content found in Word file")
                    return None

                print(f"Successfully extracted {len(text)} characters from Word file")
                return text
            except Exception as doc_error:
                print(f"Error processing Word file: {str(doc_error)}")
                return None
        elif ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                print(f"Successfully extracted {len(text)} characters from TXT")
                return text
    except Exception as e:
        print(f"Error extracting text from {file_path}: {str(e)}")
        return None

def generate_mcqs(input_data, num_questions, difficulty="easy", is_file=False):
    difficulty_prompts = {
        "easy": "Generate simple, straightforward questions suitable for beginners. Focus on basic concepts and definitions.",
        "intermediate": "Generate moderately challenging questions that test understanding and application of concepts.",
        "hard": "Generate complex questions that test deep understanding, analysis, and synthesis of concepts."
    }

    content = input_data['content']
    word_count = len(content.split())
    print(f"Input word count: {word_count}")

    MAX_CHUNK_SIZE = 3000
    MIN_CHUNK_SIZE = 500

    def split_into_chunks(text, max_size=MAX_CHUNK_SIZE):
        chunks = []
        current_chunk = ""

        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if len(current_chunk) + len(para) < max_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())
                    current_chunk = para + "\n\n"
                else:
                    sentences = [s + '.' for s in para.split('. ') if s]
                    temp = ""
                    for sent in sentences:
                        if len(temp) + len(sent) < max_size:
                            temp += sent + " "
                        else:
                            if temp.strip():
                                chunks.append(temp.strip())
                            temp = sent + " "
                    if temp.strip():
                        chunks.append(temp.strip())
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        return [c for c in chunks if len(c) >= MIN_CHUNK_SIZE]

    def summarize_chunk(chunk_text):
        prompt = f"""
        Summarize the following text concisely while preserving key facts, concepts, names, and data.
        Keep it under 200 words. Do not add opinions or external info.

        Text:
        {chunk_text}
        """
        try:
            response = model.generate_content(prompt)
            summary = response.text.strip()
            print(f"Chunk summarized to {len(summary)} chars")
            return summary
        except Exception as e:
            print(f"⚠️ Summarization failed: {e}. Using truncated chunk.")
            return chunk_text[:800]

    def generate_questions_from_summary(summary, target_num, difficulty_level):
        prompt = f"""
        Generate {target_num} multiple-choice questions from the following summarized content:
        {summary}

        {difficulty_prompts[difficulty_level]}

        Instructions:
        - Use ONLY information from the summary.
        - Ensure questions are diverse and cover different parts.
        - Each question must have exactly 4 options (A, B, C, D) and one correct answer.
        - Format exactly as:
        ## MCQ
        Question: [question text]
        A) [option A]
        B) [option B]
        C) [option C]
        D) [option D]
        Correct Answer: [letter]

        Do not add extra text or markdown.
        """
        try:
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            print(f"⚠️ MCQ generation failed for chunk: {e}")
            return None

    if word_count > 1500:
        print("🚀 Large document detected. Activating chunking pipeline...")
        chunks = split_into_chunks(content)
        print(f"📄 Split into {len(chunks)} chunks.")

        if len(chunks) == 0:
            return "Error: No valid content could be extracted."

        all_mcq_texts = []
        base_questions_per_chunk = max(1, num_questions // len(chunks))
        remainder = num_questions % len(chunks)

        for i, chunk in enumerate(chunks):
            print(f"🧠 Processing chunk {i+1}/{len(chunks)}...")
            summary = summarize_chunk(chunk)
            q_count = base_questions_per_chunk + (1 if i < remainder else 0)
            mcqs_text = generate_questions_from_summary(summary, q_count, difficulty)

            if mcqs_text:
                all_mcq_texts.append(mcqs_text)
            else:
                print(f"❌ Skipping chunk {i+1} due to generation failure.")

        final_mcqs = "\n\n".join(all_mcq_texts)
        print(f"✅ Generated MCQs from {len(all_mcq_texts)} chunks.")
        return final_mcqs

    else:
        if is_file:
            prompt = f"""
            Generate {num_questions} multiple-choice questions from the following text content:
            {content}

            {difficulty_prompts[difficulty]}

            Additional instructions for file-based questions:
            1. Use ONLY the information from the provided text content
            2. Focus on key concepts and important details from the text
            3. Include questions about specific facts, figures, or data mentioned in the text
            4. Create questions that test comprehension of the main ideas and supporting details
            5. Ensure questions are directly related to the content in the file
            6. Include at least one question about any tables, lists, or structured data if present
            7. Cover a broad range of topics from the text
            8. Ensure questions are well-distributed across different sections of the content

            Format exactly like this for each question:
            ## MCQ
            Question: [question text]
            A) [option A]
            B) [option B]
            C) [option C]
            D) [option D]
            Correct Answer: [letter of correct option]

            Ensure each question has exactly 4 options and one correct answer.
            """
        else:
            prompt = f"""
            Generate {num_questions} multiple-choice questions from the following topic:
            {input_data['content']}

            {difficulty_prompts[difficulty]}

            Additional instructions for topic-based questions:
            1. Create questions that cover the main aspects of the topic
            2. Include both theoretical and practical questions
            3. Ensure questions are relevant to the given topic
            4. Create a good mix of definition, concept, and application questions
            5. Cover fundamental concepts and advanced aspects of the topic
            6. Include questions about key terminology and important principles
            7. Ensure questions test both understanding and application
            8. Create questions that help build a comprehensive understanding of the topic

            Format exactly like this for each question:
            ## MCQ
            Question: [question text]
            A) [option A]
            B) [option B]
            C) [option C]
            D) [option D]
            Correct Answer: [letter of correct option]

            Ensure each question has exactly 4 options and one correct answer.
            """

        try:
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            print(f"Error generating MCQs: {e}")
            return None

def generate_question_paper(content_text, counts_by_marks):
    try:
        sections_prompt = []
        if counts_by_marks.get(2, 0) > 0:
            sections_prompt.append(f"- {counts_by_marks[2]} questions worth 2 marks each focused on concise facts or definitions.")
        if counts_by_marks.get(3, 0) > 0:
            sections_prompt.append(f"- {counts_by_marks[3]} questions worth 3 marks each that require short explanations.")
        if counts_by_marks.get(5, 0) > 0:
            sections_prompt.append(f"- {counts_by_marks[5]} questions worth 5 marks each requiring detailed understanding and multi-point answers.")
        if counts_by_marks.get(10, 0) > 0:
            sections_prompt.append(f"- {counts_by_marks[10]} questions worth 10 marks each demanding comprehensive, structured answers.")

        prompt = f"""
        You are an expert educator. Create an exam-style question paper from the content below.

        Content:
        {content_text}

        Create the following sections exactly, with numbered questions under each section:
        {chr(10).join(sections_prompt)}

        Rules:
        1. Do NOT include answers.
        2. Ensure questions are derived ONLY from the given content.
        3. Avoid duplication; cover a broad range of topics.
        4. Each question must clearly match its mark weight by depth.
        5. Format strictly as plain text with clear section headings: "2 Marks", "3 Marks", "5 Marks", "10 Marks".
        6. Under each heading, number the questions starting from 1.
        7. Keep language clear and unambiguous.
        """

        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print(f"Error generating question paper: {e}")
        return None

@app.route('/')
def index():
    if request.args.get('clear') == 'true':
        session.clear()
    return render_template('index.html')

@app.route('/home')
def home():
    session.clear()
    return redirect(url_for('index'))

@app.route('/generate', methods=['POST'])
def generate():
    session.clear()

    file = request.files.get('file')
    topic = request.form.get('topic')
    difficulty = request.form.get('difficulty', 'easy')

    print(f"File received: {file}")
    print(f"Topic received: {topic}")
    print(f"Request files: {request.files}")
    print(f"Request form: {request.form}")

    try:
        num_questions = int(request.form['num_questions'])
        print(f"Generating {num_questions} questions")
    except ValueError:
        return "Please enter a valid number of questions."

    if not file and topic:
        print("Processing topic input only")
        input_data = {
            'content': topic,
            'is_file': False
        }

    elif file and allowed_file(file.filename):
        print("Processing file upload only")
        try:
            filename = secure_filename(file.filename)
            timestamp = str(int(time.time()))
            unique_filename = f"{timestamp}_{filename}"
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            print(f"Saving file to: {file_path}")

            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

            file.save(file_path)
            print(f"File saved successfully")

            file_content = extract_text_from_file(file_path)
            if not file_content:
                print("Failed to extract text from file")

                try:
                    os.remove(file_path)
                except:
                    pass
                return "Error: Could not extract text from the uploaded file. Please try again."

            print(f"Extracted text length: {len(file_content)}")

            input_data = {
                'content': file_content,
                'is_file': True
            }

            session['current_file'] = file_path
            session['difficulty'] = difficulty

        except Exception as e:
            print(f"Error processing file: {str(e)}")
            return f"Error processing file: {str(e)}"

    else:
        print(f"Invalid input: file={file}, topic={topic}")
        return "Please provide either a file or a topic."

    mcqs = generate_mcqs(input_data, num_questions, difficulty)
    if not mcqs:
        print("Failed to generate MCQs")
        return "Failed to generate MCQs. Please try again."

    print(f"Generated MCQs: {mcqs[:200]}...")

    questions = []
    for mcq in mcqs.split("## MCQ"):
        if not mcq.strip():
            continue

        lines = [line.strip() for line in mcq.split('\n') if line.strip()]
        if len(lines) < 6:
            continue

        question = lines[0].replace("Question:", "").strip()
        options = [
            lines[1][3:].strip(),
            lines[2][3:].strip(),
            lines[3][3:].strip(),
            lines[4][3:].strip()
        ]
        correct_answer = lines[5].replace("Correct Answer:", "").strip()

        questions.append({
            'question': question,
            'options': options,
            'correct_answer': options[ord(correct_answer.upper()) - 65]
        })

    print(f"Processed {len(questions)} questions")

    if not questions:
        print("No valid questions generated")
        return "No valid questions could be generated. Please try again."

    # ✅ SAVE QUESTIONS TO SERVER-SIDE FILE
    quiz_id = str(uuid.uuid4())
    quiz_file_path = os.path.join(QUIZ_STORAGE_DIR, f"{quiz_id}.json")

    with open(quiz_file_path, 'w', encoding='utf-8') as f:
        json.dump({
            'questions': questions,
            'original_input_text': input_data['content'] if input_data['is_file'] else None,
            'difficulty': difficulty,
            'num_questions': num_questions,
            'generated_at': time.time()
        }, f, ensure_ascii=False, indent=2)

    # ✅ ONLY STORE QUIZ ID IN SESSION (not the whole list!)
    session['quiz_id'] = quiz_id

    return redirect(url_for('quiz'))

@app.route('/generate_paper', methods=['POST'])
def generate_paper():
    try:
        session.clear()
        file = request.files.get('file_paper')
        topic = request.form.get('topic_paper')
        try:
            counts_by_marks = {
                2: int(request.form.get('count_2', '0') or '0'),
                3: int(request.form.get('count_3', '0') or '0'),
                5: int(request.form.get('count_5', '0') or '0'),
                10: int(request.form.get('count_10', '0') or '0'),
            }
        except ValueError:
            return "Please enter valid counts for question marks."

        if all(v == 0 for v in counts_by_marks.values()):
            return "Please request at least one question."

        content_text = None
        if file and allowed_file(file.filename):
            try:
                filename = secure_filename(file.filename)
                timestamp = str(int(time.time()))
                unique_filename = f"{timestamp}_{filename}"
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
                os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
                file.save(file_path)
                content_text = extract_text_from_file(file_path)
                session['current_file'] = file_path
            except Exception as e:
                print(f"Error processing paper file: {str(e)}")
                return f"Error processing file: {str(e)}"
        elif topic:
            content_text = topic
        else:
            return "Please provide either a file or a topic."

        paper_text = generate_question_paper(content_text, counts_by_marks)
        if not paper_text:
            return "Failed to generate question paper. Please try again."

        session['question_paper'] = paper_text
        session['question_paper_counts'] = counts_by_marks
        # Generate answer key aligned to the question paper
        try:
            answer_key_prompt = f"""
            Using ONLY the content below and respecting the exact numbering and sections
            of the given Question Paper, generate a concise Answer Key.

            Content:
            {content_text}

            Question Paper:
            {paper_text}

            Requirements:
            - Follow the same headings (2 Marks, 3 Marks, 5 Marks, 10 Marks)
            - Number answers exactly matching the questions
            - Keep answers concise and correct; no extra commentary
            - If a question is open-ended, provide key points or a model answer
            - Plain text only
            """
            response = model.generate_content(answer_key_prompt)
            session['answer_key'] = response.text
        except Exception as e:
            print(f"Error generating answer key: {e}")
            session['answer_key'] = None
        return redirect(url_for('paper'))
    except Exception as e:
        print(f"Error in generate_paper: {str(e)}")
        return "An error occurred while generating the question paper.", 500

@app.route('/paper')
def paper():
    if 'question_paper' not in session:
        return redirect(url_for('index'))
    return render_template('paper.html', paper_text=session['question_paper'], counts=session.get('question_paper_counts', {}), has_answer_key=session.get('answer_key') is not None)

@app.route('/download_paper_docx')
def download_paper_docx():
    try:
        paper_text = session.get('question_paper')
        if not paper_text:
            return redirect(url_for('index'))

        doc = DocxDocument()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)

        for block in paper_text.split('\n\n'):
            p = doc.add_paragraph()
            for line in block.split('\n'):
                p.add_run(line)
                p.add_run('\n')
        filename = f"question_paper_{int(time.time())}.docx"
        out_path = os.path.join(app.config['RESULTS_FOLDER'], filename)
        doc.save(out_path)
        return send_file(out_path, as_attachment=True)
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        return "Failed to prepare DOCX.", 500

@app.route('/download_answer_key_docx')
def download_answer_key_docx():
    try:
        answer_key = session.get('answer_key')
        if not answer_key:
            return redirect(url_for('paper'))

        doc = DocxDocument()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)

        title = doc.add_paragraph()
        title_run = title.add_run('Answer Key')
        title_run.bold = True

        for block in answer_key.split('\n\n'):
            p = doc.add_paragraph()
            for line in block.split('\n'):
                p.add_run(line)
                p.add_run('\n')
        filename = f"answer_key_{int(time.time())}.docx"
        out_path = os.path.join(app.config['RESULTS_FOLDER'], filename)
        doc.save(out_path)
        return send_file(out_path, as_attachment=True)
    except Exception as e:
        print(f"Error generating Answer Key DOCX: {e}")
        return "Failed to prepare Answer Key DOCX.", 500

@app.route('/chat')
def chat():
    return render_template('chat.html')

@app.route('/chat_upload', methods=['POST'])
def chat_upload():
    try:
        file = request.files.get('file')
        if not file or not allowed_file(file.filename):
            return redirect(url_for('chat'))

        filename = secure_filename(file.filename)
        timestamp = str(int(time.time()))
        unique_filename = f"{timestamp}_{filename}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        file.save(file_path)

        # Validate extraction to ensure it's a usable context
        _ = extract_text_from_file(file_path)
        session['current_file'] = file_path
        return redirect(url_for('chat'))
    except Exception:
        return redirect(url_for('chat'))

@app.route('/chat_from_doc', methods=['POST'])
def chat_from_doc():
    try:
        data = request.get_json()
        user_question = data.get('question', '').strip() if data else ''
        if not user_question:
            return jsonify({ 'status': 'error', 'error': 'Question is required.' }), 400

        context_text = None

        current_file = session.get('current_file')
        if current_file and os.path.exists(current_file):
            extracted = extract_text_from_file(current_file)
            if extracted:
                context_text = extracted

        if not context_text:
            quiz_id = session.get('quiz_id')
            if quiz_id:
                try:
                    quiz_file_path = os.path.join(QUIZ_STORAGE_DIR, f"{quiz_id}.json")
                    if os.path.exists(quiz_file_path):
                        with open(quiz_file_path, 'r', encoding='utf-8') as f:
                            quiz_data = json.load(f)
                            context_text = quiz_data.get('original_input_text') or quiz_data.get('content') or ''
                except Exception:
                    pass

        if not context_text:
            # Fallback to last generated paper text if available
            paper_text = session.get('question_paper')
            if paper_text:
                context_text = paper_text

        if not context_text:
            return jsonify({ 'status': 'error', 'error': 'No document context available. Upload a PDF/DOCX/TXT or generate content first.' }), 400

        prompt = f"""
        You are a helpful tutor. Answer the student's question ONLY using the content below.
        If the answer is not present in the content, say you don't have enough information.
        Be clear and concise; use step-by-step explanation when beneficial.

        Document Content:
        {context_text}

        Question:
        {user_question}
        """
        response = model.generate_content(prompt)
        return jsonify({ 'status': 'success', 'answer': response.text })
    except Exception as e:
        return jsonify({ 'status': 'error', 'error': str(e) }), 500

@app.route('/quiz', methods=['GET', 'POST'])
def quiz():
    # ✅ LOAD QUESTIONS FROM FILE USING quiz_id
    quiz_id = session.get('quiz_id')
    if not quiz_id:
        print("No quiz_id in session → redirecting to index")
        return redirect(url_for('index'))

    quiz_file_path = os.path.join(QUIZ_STORAGE_DIR, f"{quiz_id}.json")
    if not os.path.exists(quiz_file_path):
        print(f"Quiz file {quiz_file_path} not found → redirecting to index")
        return redirect(url_for('index'))

    try:
        with open(quiz_file_path, 'r', encoding='utf-8') as f:
            quiz_data = json.load(f)
            questions = quiz_data['questions']
    except Exception as e:
        print(f"Error loading quiz data: {e}")
        return redirect(url_for('index'))

    if request.method == 'POST':
        try:
            user_answers = []
            score = 0

            for i, q in enumerate(questions):
                user_answer = request.form.get(f"question_{i}")
                if not user_answer:
                    user_answer = "No answer selected"

                correct = user_answer == q['correct_answer']

                if correct:
                    score += 1

                prompt = f"""
                Explain why '{q['correct_answer']}' is the correct answer to:
                '{q['question']}'

                Provide a clear, concise explanation in 2-3 sentences.
                """
                try:
                    response = model.generate_content(prompt)
                    explanation = response.text
                except Exception as e:
                    print(f"Error generating explanation: {str(e)}")
                    explanation = "Explanation not available."

                user_answers.append({
                    'question': q['question'],
                    'options': q['options'],
                    'correct_answer': q['correct_answer'],
                    'user_answer': user_answer,
                    'is_correct': correct,
                    'explanation': explanation
                })

            # ✅ Store only small data in session
            session['user_answers'] = user_answers
            session['score'] = score
            session['total'] = len(questions)

            create_pdf(user_answers, score, len(questions))

            percentage = (score / len(questions) * 100) if len(questions) > 0 else 0

            return render_template('scoreboard.html',
                                 user_answers=user_answers,
                                 score=score,
                                 total=len(questions),
                                 percentage=percentage,
                                 chr=chr,
                                 enumerate=enumerate)

        except Exception as e:
            print(f"Error processing quiz submission: {str(e)}")
            return redirect(url_for('index'))

    return render_template('quiz.html', questions=questions, enumerate=enumerate)

@app.route('/scoreboard')
def scoreboard():
    try:
        score = session.get('score', 0)
        total = session.get('total', 0)
        user_answers = session.get('user_answers', [])

        if not user_answers or total == 0:
            return redirect(url_for('index'))

        percentage = (score / total * 100) if total > 0 else 0

        try:
            prompt = f"""
            Create comprehensive study notes based on the following quiz content:

            Quiz Questions and Answers:
            {chr(10).join([f"Q: {q['question']}{chr(10)}A: {q['correct_answer']}{chr(10)}E: {q['explanation']}{chr(10)}" for q in user_answers])}

            Create detailed study notes that:
            1. Start with a clear introduction explaining the main topic
            2. Break down each concept from the quiz questions into detailed explanations
            3. Include:
               - Key definitions and terms
               - Important concepts and their explanations
               - Examples and applications
               - Practice tips and key takeaways
            4. Organize the content into clear sections:
               - Introduction
               - Main Concepts
               - Detailed Explanations
               - Examples
               - Summary
               - Practice Tips
            5. Make the content:
               - Easy to understand
               - Well-structured
               - Educational and informative
               - Helpful for learning and revision
            6. Use bullet points for important points
            7. Include all explanations from the quiz
            8. Add additional relevant information to help understand the concepts
            9. Make it suitable for both beginners and advanced learners
            10. Ensure it's comprehensive but concise

            Format the output in a clear, readable way with proper sections and bullet points.
            Do not use any markdown symbols like ** or ##. Use simple text formatting with clear section headers.
            Make it suitable for creating a PDF study guide.
            """

            response = model.generate_content(prompt)
            notes_content = response.text

            pdf = FPDF()
            pdf.add_page()

            pdf.set_font("Arial", "B", 16)
            pdf.cell(200, 10, txt="Study Notes", ln=1, align='C')
            pdf.ln(10)

            pdf.set_font("Arial", "", 12)

            sections = notes_content.split('\n\n')
            for section in sections:
                lines = section.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        pdf.ln(5)
                    elif line.endswith(':'):
                        pdf.set_font('Arial', 'B', 12)
                        pdf.cell(0, 10, txt=line, ln=1)
                        pdf.set_font('Arial', '', 12)
                    elif line.startswith('•'):
                        pdf.set_x(10)
                        pdf.multi_cell(0, 10, txt=line[1:].strip())
                    else:
                        pdf.multi_cell(0, 10, txt=line)
                pdf.ln(5)

            timestamp = str(int(time.time()))
            notes_filename = f'study_notes_{timestamp}.pdf'
            notes_path = os.path.join(app.config['RESULTS_FOLDER'], notes_filename)
            pdf.output(notes_path)

            session['current_notes_file'] = notes_filename
            print(f"Notes generated successfully: {notes_filename}")

        except Exception as e:
            print(f"Error generating notes: {str(e)}")
            session['current_notes_file'] = None

        return render_template('scoreboard.html',
                             user_answers=user_answers,
                             score=score,
                             total=total,
                             percentage=percentage,
                             chr=chr,
                             enumerate=enumerate)
    except Exception as e:
        print(f"Error in scoreboard: {str(e)}")
        return redirect(url_for('index'))

@app.route('/get_reasoning', methods=['POST'])
def get_reasoning():
    try:
        data = request.get_json()
        question = data.get('question')
        correct_answer = data.get('correct_answer')

        prompt = f"""
        Explain why '{correct_answer}' is the correct answer to:
        '{question}'

        Provide a clear, concise explanation in 2-3 sentences.
        """

        response = model.generate_content(prompt)
        return jsonify({
            'reasoning': response.text,
            'status': 'success'
        })
    except Exception as e:
        return jsonify({
            'error': str(e),
            'status': 'error'
        }), 500

@app.route('/generate_notes', methods=['POST'])
def generate_notes():
    try:
        user_answers = session.get('user_answers', [])
        if not user_answers:
            return jsonify({
                'status': 'error',
                'error': 'No quiz data available.'
            }), 400

        questions_text = ""
        for q in user_answers:
            questions_text += f"Q: {q['question']}\n"
            questions_text += f"A: {q['correct_answer']}\n"
            questions_text += f"E: {q['explanation']}\n\n"

        prompt = f"""
        Create comprehensive study notes based on the following quiz content:

        Quiz Questions and Answers:
        {questions_text}

        Create detailed study notes that:
        1. Start with a clear introduction explaining the main topic
        2. Break down each concept from the quiz questions into detailed explanations
        3. Include:
           - Key definitions and terms
           - Important concepts and their explanations
           - Examples and applications
           - Practice tips and key takeaways
        4. Organize the content into clear sections:
           - Introduction
           - Main Concepts
           - Detailed Explanations
           - Examples
           - Summary
           - Practice Tips
        5. Make the content:
           - Easy to understand
           - Well-structured
           - Educational and informative
           - Helpful for learning and revision
        6. Use bullet points for important points
        7. Include all explanations from the quiz
        8. Add additional relevant information to help understand the concepts
        9. Make it suitable for both beginners and advanced learners
        10. Ensure it's comprehensive but concise

        Format the output in a clear, readable way with proper sections and bullet points.
        Do not use any markdown symbols like ** or ##. Use simple text formatting with clear section headers.
        Make it suitable for creating a PDF study guide.
        """

        response = model.generate_content(prompt)
        notes_content = response.text

        pdf = FPDF()
        pdf.add_page()

        pdf.set_font("Arial", "B", 16)
        pdf.cell(200, 10, txt="Study Notes", ln=1, align='C')
        pdf.ln(10)

        pdf.set_font("Arial", "", 12)

        sections = notes_content.split('\n\n')
        for section in sections:
            lines = section.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    pdf.ln(5)
                elif line.endswith(':'):
                    pdf.set_font('Arial', 'B', 12)
                    pdf.cell(0, 10, txt=line, ln=1)
                    pdf.set_font('Arial', '', 12)
                elif line.startswith('•'):
                    pdf.set_x(10)
                    pdf.multi_cell(0, 10, txt=line[1:].strip())
                else:
                    pdf.multi_cell(0, 10, txt=line)
            pdf.ln(5)

        timestamp = str(int(time.time()))
        notes_filename = f'study_notes_{timestamp}.pdf'
        notes_path = os.path.join(app.config['RESULTS_FOLDER'], notes_filename)
        pdf.output(notes_path)

        return jsonify({
            'status': 'success',
            'notes_path': notes_filename
        })

    except Exception as e:
        print(f"Error generating notes: {str(e)}")
        return jsonify({
            'status': 'error',
            'error': str(e)
        }), 500

def create_pdf(user_answers, score, total):
    pdf = FPDF()
    pdf.add_page()

    pdf.set_font("Helvetica", size=12)

    pdf.cell(0, 10, f"Score: {score}/{total}", ln=True)
    pdf.ln(10)

    for ans in user_answers:
        pdf.multi_cell(0, 10, f"Q: {ans['question']}")

        for i, option in enumerate(ans['options']):
            status = ""
            if option == ans['correct_answer']:
                status = "(Correct Answer)"
            elif option == ans['user_answer']:
                status = "(Your Answer)"
            pdf.multi_cell(0, 10, f"{chr(65+i)}) {option} {status}")

        pdf.ln(10)

    pdf_path = os.path.join(app.config['RESULTS_FOLDER'], 'results.pdf')
    pdf.output(pdf_path)
    return pdf_path

@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_path = os.path.join(app.config['RESULTS_FOLDER'], filename)
        print(f"Attempting to download file: {file_path}")

        if not os.path.exists(file_path):
            print(f"File not found at: {file_path}")
            return "File not found", 404

        print(f"File found, sending download")
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        print(f"Error downloading file: {str(e)}")
        return str(e), 500

if __name__ == "__main__":
    app.run(debug=True)
